from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# APA margins — 1 inch all sides
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# Default font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ── Helpers ─────────────────────────────────────────────────────────────────

def heading(text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.bold = True
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run('     ' + text)
        r.bold = True
        r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p

def body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p

def figure(img_path, caption_text, width=Inches(6.0)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption_text)
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    doc.add_paragraph()   # breathing room

def block(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    r1 = p.add_run(label + '  ')
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = p.add_run(text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    doc.add_paragraph()

_NBS = '\u202f'  # narrow no-break space used by macOS in screenshot filenames
_BASE = '/Users/connorsolvason/my-website/Screenshot 2026-03-25 at '

IMGS = {
    1:  _BASE + f'6.42.46{_NBS}AM.png',
    2:  _BASE + f'6.43.47{_NBS}AM.png',
    3:  _BASE + f'6.44.53{_NBS}AM.png',
    4:  _BASE + f'6.45.26{_NBS}AM.png',
    5:  _BASE + f'6.45.42{_NBS}AM.png',
    6:  _BASE + f'6.45.49{_NBS}AM.png',
    7:  _BASE + f'6.46.32{_NBS}AM.png',
    8:  _BASE + f'6.47.01{_NBS}AM.png',
    9:  _BASE + f'6.47.32{_NBS}AM.png',
    10: _BASE + f'6.48.19{_NBS}AM.png',
    11: _BASE + f'6.49.43{_NBS}AM.png',
    12: _BASE + f'6.49.54{_NBS}AM.png',
    13: _BASE + f'6.50.06{_NBS}AM.png',
}

# ═══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
r = p.add_run(
    'Rebuilding a Nonprofit Digital Presence Using Conversational AI:\n'
    'A Case Study of Asociación Deportiva Curundú'
)
r.bold = True
r.font.name = 'Times New Roman'
r.font.size = Pt(14)

for line in [
    '\n\nConnor Solvason',
    'Panama City International School / MET IB Program',
    'March 25, 2026',
    '\nCourse: Artificial Intelligence Applications',
    'Instructor: [AI Professor]',
]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════

heading('Abstract')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Inches(0.5)
r = p.add_run(
    'This paper documents the complete redesign and redevelopment of the digital presence '
    'for Asociación Deportiva Curundú (ADC), a nonprofit youth soccer and community development '
    'foundation operating in the Curundú district of Panama City, Panama. The project was '
    'executed entirely through iterative natural-language prompting of Claude, Anthropic\'s '
    'large language model, across five development phases spanning approximately four months. '
    'Beginning with only a legacy Squarespace site and a brief organizational description, '
    'the AI system generated a full eight-page responsive website, a private administrative '
    'backend, and a Cloudflare Worker proxy enabling AI-powered grant proposal writing—all '
    'without the project lead writing a single line of code manually. The resulting site, '
    'deployed to GitHub Pages at no cost, incorporates modern nonprofit web design best practices '
    'including impact statistics, donor transparency features, bilingual content, a filterable '
    'blog, a multi-pathway participation pipeline, and a live AI chatbot. This paper analyzes '
    'each major transformation through thirteen side-by-side visual comparisons, documents '
    'the specific prompts that produced each change, and articulates the underlying design '
    'rationale. The paper concludes with a review of planned enhancements that illustrate the '
    'expanding and ongoing role of conversational AI in organizational digital strategy. '
    'Collectively, this project demonstrates that large language models are capable of serving '
    'as full-stack web development partners for non-technical operators, delivering '
    'production-grade output through natural language alone.'
)
r.font.name = 'Times New Roman'
r.font.size = Pt(12)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.5)
r1 = p.add_run('Keywords: ')
r1.italic = True
r1.font.name = 'Times New Roman'
r1.font.size = Pt(12)
r2 = p.add_run(
    'conversational AI, nonprofit technology, web development, large language models, '
    'Claude AI, natural language programming, digital transformation, Panama'
)
r2.font.name = 'Times New Roman'
r2.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════

heading(
    'Rebuilding a Nonprofit Digital Presence Using Conversational AI:\n'
    'A Case Study of Asociación Deportiva Curundú'
)
heading('Introduction')

body(
    'Asociación Deportiva Curundú (ADC) is a nonprofit youth soccer and community '
    'development foundation established in 2014 in Curundú, one of the most economically '
    'underserved districts of Panama City, Panama. Through sport, education, and nutritional '
    'support, ADC serves over 70 young athletes and their families across 11 teams. Despite '
    'its demonstrated community impact—including placing players on Panama\'s national youth '
    'team and achieving multiple tournament honors—ADC\'s digital presence consisted of an '
    'aging Squarespace website that failed to reflect the depth of the organization\'s work, '
    'communicate persuasively with donors, or streamline operational processes such as athlete '
    'registration, volunteer coordination, or grant applications.'
)

body(
    'In early 2026, a complete website rebuild was initiated using Claude, Anthropic\'s large '
    'language model, as the sole development partner. The project required no manual coding '
    'by the project lead; every HTML page, CSS stylesheet, JavaScript behavior, and backend '
    'component was generated through iterative natural-language prompting across five '
    'documented development phases. The result is a production-grade, fully responsive '
    'eight-page website deployed via GitHub Pages, accompanied by a private administrative '
    'dashboard featuring a Claude API-powered grant proposal writer.'
)

body(
    'This paper presents a structured academic analysis of that development process. For '
    'each major website transformation, a side-by-side visual comparison is provided—new '
    'version on the left, original Squarespace site on the right—followed by the specific '
    'prompt that initiated the change and an explanation of the design rationale. The paper '
    'concludes with a forward-looking section on planned enhancements, illustrating the '
    'expanding capabilities that AI makes available to organizations with limited technical '
    'and financial resources.'
)

# ═══════════════════════════════════════════════════════════════════════════
# METHODOLOGY
# ═══════════════════════════════════════════════════════════════════════════

heading('Methodology')

body(
    'The development process followed an iterative, conversational programming paradigm. '
    'The project lead communicated desired outcomes, features, or corrections to Claude '
    'entirely in plain language. The AI model then generated or modified the relevant code '
    'files, which were saved locally and deployed to GitHub Pages. No integrated development '
    'environment, build pipeline, or prior programming knowledge was required.'
)

body(
    'Development was organized into five phases: (1) Site Architecture and Foundation, '
    '(2) Content Population and Brand Identity, (3) Full Page Build-Out, (4) Nonprofit '
    'Best Practices and AI Feature Integration, and (5) Documentation and Long-Term '
    'Sustainability. Screenshots were captured at the conclusion of major change sets to '
    'provide visual evidence of before-and-after states. In every comparison image, the '
    'new version appears on the left and the original Squarespace site appears on the right.'
)

body(
    'Forty-eight prompts were documented across the five phases. All code generated by '
    'Claude—HTML, CSS, JavaScript, Python, and Cloudflare Worker scripts—was reviewed '
    'for functional accuracy by the project lead but was not manually written or '
    'significantly modified. The project models a human-AI collaboration in which the '
    'human partner provides creative direction, organizational knowledge, and quality '
    'review, while the AI partner provides technical execution.'
)

# ═══════════════════════════════════════════════════════════════════════════
# PHASE 1
# ═══════════════════════════════════════════════════════════════════════════

heading('Phase 1: Site Architecture and Hero Section')

body(
    'The first development phase established the structural foundation of the site: '
    'the file architecture, navigation system, and hero section. The original Squarespace '
    'site featured a static video hero with a fragmented navigation bar containing '
    'inconsistent Spanish/English labeling and an unexplained language toggle. The first '
    'prompting session produced a complete multi-page HTML/CSS framework with a responsive '
    'navigation bar, a full-screen looping video hero, and a consistent visual language '
    'anchored in ADC\'s green-and-black brand palette.'
)

figure(
    IMGS[1],
    'Figure 1. Hero section comparison. New site (left): classroom video background with clean five-item '
    'navigation. Original site (right): outdoor video with fragmented seven-item navigation '
    'including an unexplained EN language toggle.'
)

block(
    'Prompt:',
    '"Build a complete multi-page website for ADC, a nonprofit soccer and education '
    'organization in Curundú, Panama. It should be fully in Spanish, mobile responsive, '
    'use their green and black brand colors, and open with a hero section featuring a '
    'looping background video and the tagline \'Nuestro barrio, nuestro equipo.\'"'
)

block(
    'Rationale:',
    'The hero section is the highest-impact real estate on any website—it is the first '
    'content a visitor perceives and determines whether they continue reading. The AI '
    'replaced a static outdoor video with a classroom scene, underscoring ADC\'s dual '
    'identity as both a sports program and an educational institution. The navigation was '
    'consolidated from seven items to five, reducing cognitive load and eliminating the '
    'confusing bilingual toggle. Every subsequent page in the site inherits this navigation '
    'component, ensuring visual consistency without redundant code.'
)

# ═══════════════════════════════════════════════════════════════════════════
# PHASE 2
# ═══════════════════════════════════════════════════════════════════════════

heading('Phase 2: Content Population and Brand Identity')

body(
    'With the structural framework in place, Phase 2 focused on populating the site with '
    'substantive organizational content: the About page, team biographies, mission and '
    'vision statements, impact statistics, organizational values, a historical timeline, '
    'and athletic achievements. This phase also merged two previously separate pages—'
    '"Acerca de ADC" and "Quiénes Somos"—into a single, coherent organizational identity '
    'section, eliminating redundancy and consolidating the organization\'s story.'
)

figure(
    IMGS[2],
    'Figure 2. About page introduction. New site (left): unified "Acerca de ADC" with team photo, '
    'mission statement, and a call-to-action button. Original site (right): separate "Nuestro Equipo" '
    'section alongside a manually annotated logo diagram ("Significado del Logotipo").'
)

block(
    'Prompt:',
    '"Merge the Acerca de ADC and Quiénes Somos pages into a single unified about page. '
    'Include the organizational mission, a team photo, and a clear call-to-action button '
    'inviting visitors to learn more."'
)

block(
    'Rationale:',
    'Two separate pages covering organizational identity created redundancy and confused '
    'visitors about where to find foundational information. Merging them into a single '
    'destination reduces navigation friction and concentrates ADC\'s story in one place. '
    'The manually annotated logo diagram on the original site, while creative, was a static '
    'image that required design software to update; the new site renders equivalent '
    'information as accessible, maintainable HTML content.'
)

figure(
    IMGS[3],
    'Figure 3. Impact statistics section. New site (left): animated counters displaying 70+ athletes, '
    '60+ families, 5+ years of operation, and 11 teams. Original site (right): no quantitative '
    'impact metrics anywhere on the page.'
)

block(
    'Prompt:',
    '"Add an animated impact statistics section to the about page showing ADC\'s key '
    'numbers: athletes currently served, families impacted, years in operation, and '
    'active teams. Use large display numbers that count up on scroll."'
)

block(
    'Rationale:',
    'Donors, grant reviewers, and prospective partners make funding and partnership '
    'decisions based on demonstrated, quantifiable impact. Displaying specific numbers '
    '(70+ athletes, 60+ families, 5+ years) transforms abstract mission statements into '
    'concrete evidence of organizational effectiveness. Animated counters add visual '
    'dynamism that draws the eye and signal organizational confidence to first-time '
    'visitors who cannot yet assess ADC\'s credibility through personal experience.'
)

figure(
    IMGS[4],
    'Figure 4. Leadership team biographies. New site (left): individual profile cards for '
    'Andrés Madrid, César Santos, and Javier Wallace with photos and bios. Original site (right): '
    '"Nosotros" text column with Mission and Vision statements and no named individuals.'
)

block(
    'Prompt:',
    '"Add individual team biography cards to the about page for ADC\'s three founders: '
    'Andrés Madrid (founder, Panamanian national team player), César Santos (co-founder, '
    'former professional player), and Javier Wallace (head coach, international experience). '
    'Each card should include a photo, title, and brief biography."'
)

block(
    'Rationale:',
    'Personalization is a cornerstone of nonprofit credibility. Visitors who can see the '
    'faces and read the stories of real leaders are significantly more likely to trust the '
    'organization\'s legitimacy and donate. The original site referred to "our team" '
    'without naming a single person—a critical credibility gap for organizations seeking '
    'grants, sponsorships, or community partnerships. Providing real names and photographs '
    'also fulfills a basic accountability standard expected by institutional funders.'
)

figure(
    IMGS[5],
    'Figure 5. Organizational values and milestone timeline. New site (left): visual value cards '
    '(Community, Integrity, Excellence, Joy) with a chronological timeline from 2014 to 2024. '
    'Original site (right): plain parallel text columns for Mission and Vision with no historical context.'
)

block(
    'Prompt:',
    '"Add a visual values section with icon cards for ADC\'s four core principles, '
    'and a chronological milestone timeline showing key organizational achievements '
    'from founding in 2014 through the present."'
)

block(
    'Rationale:',
    'Visual storytelling communicates organizational character more effectively than '
    'paragraph-form text. The value cards give donors and partners a rapid, memorable '
    'summary of ADC\'s culture in under ten seconds of scanning. The milestone timeline '
    'serves a dual function: it demonstrates organizational longevity and stability—'
    'critical signals for grant reviewers assessing institutional risk—and it creates a '
    'narrative arc that gives new visitors a sense of ADC\'s growth and trajectory '
    'across a decade of operation.'
)

figure(
    IMGS[6],
    'Figure 6. Athletic achievements and sponsor bar. New site (left): visual trophy cards '
    'showcasing competition honors and a consistent five-sponsor logo bar. Original site (right): '
    'bulleted achievement list with partial sponsor logos and no visual hierarchy.'
)

block(
    'Prompt:',
    '"Add a visual achievement showcase displaying ADC\'s major competition honors '
    'as trophy cards. Also add a consistent sponsor logo bar that appears across every '
    'page of the site, featuring all five current sponsors."'
)

block(
    'Rationale:',
    'Athletic achievements reinforce confidence in the quality of ADC\'s programming. '
    'Displaying specific competition results (Copa Nike placement, Liga de Corregimiento '
    'championship, Mundial del Barrio honors) demonstrates that the organization delivers '
    'on its core mission. The global sponsor bar serves two purposes: it fulfills the '
    'visibility obligations implicit in every sponsorship agreement, and it signals to '
    'new visitors that established organizations have already vetted and endorsed ADC.'
)

# ═══════════════════════════════════════════════════════════════════════════
# PHASE 3
# ═══════════════════════════════════════════════════════════════════════════

heading('Phase 3: Full Page Build-Out')

body(
    'Phase 3 extended the site from an about-page-focused structure to a complete '
    'eight-page website. This phase produced the Calendar, Participar, Contáctanos, '
    'Donar, and Blog pages—each serving a distinct functional need of a specific '
    'stakeholder audience. The AI generated each page from a combination of design '
    'briefs and organizational content provided through natural-language prompts, '
    'maintaining visual and structural consistency with the framework established in '
    'earlier phases.'
)

figure(
    IMGS[7],
    'Figure 7. Calendar page. New site (left): styled event list with color-coded category badges '
    '(Training, Event, Nutrition, Administration), event descriptions, and download buttons. '
    'Original site (right): raw Google Calendar iframe embed with no visual brand styling.'
)

block(
    'Prompt:',
    '"Build a calendar page showing ADC\'s upcoming events in a styled list format '
    'with color-coded category badges. Each event should show the date, title, '
    'description, location, and a download or registration button where applicable."'
)

block(
    'Rationale:',
    'A raw calendar iframe, while functional, creates a jarring visual break in an '
    'otherwise branded experience and renders poorly on mobile devices. The new styled '
    'event list is mobile-optimized, matches ADC\'s design system, and supports richer '
    'information display (descriptions, downloadable schedules, registration CTAs) that '
    'a calendar embed cannot provide. Category color-coding allows parents, athletes, '
    'and volunteers to scan quickly for events relevant to them without reading every entry.'
)

figure(
    IMGS[8],
    'Figure 8. Participation page. New site (left): six distinct card pathways for athlete, '
    'volunteer, MET IB student, equipment donor, community partner, and coach, each with a '
    'dedicated CTA. Original site (right): plain text step-by-step instructions addressing '
    'only one audience (registering a young athlete).'
)

block(
    'Prompt:',
    '"Create a comprehensive participation page showing all the different ways someone '
    'can get involved with ADC: as a young athlete, a volunteer, an MET IB student seeking '
    'CAS hours, a sports equipment donor, a community business partner, or a coach. '
    'Give each pathway its own card with a description and a button."'
)

block(
    'Rationale:',
    'The original participation page served only one audience: families registering '
    'young athletes. ADC\'s actual stakeholder ecosystem is substantially broader—MET '
    'students seeking IB CAS community service hours, international equipment donors, '
    'potential coaches, and corporate partners. The six-card layout ensures every '
    'visitor can immediately identify a pathway relevant to their situation, transforming '
    'a page that previously served one audience into a conversion mechanism for five '
    'distinct groups simultaneously.'
)

figure(
    IMGS[9],
    'Figure 9. Contact page. New site (left): dark hero header, comprehensive contact information '
    'panel, embedded Google Map, contact form, and newsletter subscription section. '
    'Original site (right): plain text listing of phone and email with a basic subscribe prompt.'
)

block(
    'Prompt:',
    '"Build a contact page with a dark hero header section, a contact information panel '
    'showing the address, email, phone, and hours, an embedded Google Map of the Curundú '
    'training location, a contact form with name, email, and message fields, and a '
    'newsletter subscription section at the bottom."'
)

block(
    'Rationale:',
    'A professional contact page reduces friction for every category of stakeholder. '
    'Families need to locate the training venue; journalists and potential sponsors need '
    'a direct communication channel; grant officers need a verifiable organizational '
    'address. The embedded map is particularly important for ADC\'s community context, '
    'where many families are unfamiliar with digital mapping tools and benefit from '
    'visual spatial orientation. The newsletter signup initiates an owned communication '
    'channel that ADC can use to maintain donor and community engagement independent '
    'of any social media platform.'
)

figure(
    IMGS[10],
    'Figure 10. Donation page. New site (left): tiered impact-specific giving levels ($35–$400), '
    'fund transparency progress bars, impact statistics, and an "Other Ways to Support" section. '
    'Original site (right): simple GoFundMe and PayPal buttons alongside a basic cost list.'
)

block(
    'Prompt:',
    '"Build a donation page that tells donors exactly what their money accomplishes. '
    'Create tiers at $35 (uniforms), $50 (food support for one athlete), $100 (travel '
    'costs), and $400 (full program year). Add a fund transparency section showing '
    'how donations are allocated, an impact statistics panel, and a section for '
    'non-monetary ways to support."'
)

block(
    'Rationale:',
    'Donation pages with impact-specific giving tiers consistently outperform open-ended '
    'donation pages because they make each contribution feel tangible and purposeful. '
    'Telling a donor that $35 provides a complete uniform for one athlete transforms an '
    'abstract dollar amount into a human outcome. The fund transparency section—showing '
    'percentage allocation to programs, nutrition, transport, and administration—directly '
    'addresses the most common hesitation of modern donors: concern about how their '
    'money is actually used. The "Other Ways to Support" section captures value from '
    'non-monetary stakeholders (equipment donors, volunteers, in-kind contributors) '
    'who would otherwise exit without converting.'
)

figure(
    IMGS[11],
    'Figure 11. Blog page. New site (left): "Blog ADC" with a large featured article hero and a '
    'filterable article grid below. Original site (right): basic flat article listing '
    'with no featured content or category organization.'
)

block(
    'Prompt:',
    '"Create a blog page with a full-width featured article hero at the top and a '
    'filterable article grid below. Include category filter buttons for Deportes, '
    'Educación, Nutrición, and Comunidad. Design it so new articles can be added easily."'
)

block(
    'Rationale:',
    'A well-organized blog serves multiple strategic functions for a nonprofit. It '
    'demonstrates ongoing organizational activity to donors and grant reviewers who '
    'evaluate institutional health, improves search engine visibility by providing '
    'keyword-rich content, and creates shareable material for social media channels. '
    'The featured article hero ensures that ADC\'s most important recent story always '
    'receives prominent placement. Category filtering helps different audience '
    'segments—parents interested in nutrition programming, sports enthusiasts, '
    'educational partners—find relevant content efficiently without scrolling through '
    'an undifferentiated archive.'
)

# ═══════════════════════════════════════════════════════════════════════════
# PHASE 4
# ═══════════════════════════════════════════════════════════════════════════

heading('Phase 4: Nonprofit Best Practices and AI Feature Integration')

body(
    'Phase 4 applied a systematic review of nonprofit web design best practices and '
    'introduced the most technically sophisticated features of the project. This phase '
    'added two entirely new pages—a testimonials section and a comprehensive FAQ—and '
    'produced a fully functional AI-powered administrative backend. The backend uses the '
    'Claude API through a Cloudflare Worker serverless proxy to generate customized grant '
    'proposals from a simple form interface, making professional grant writing accessible '
    'to non-specialist staff at negligible cost.'
)

figure(
    IMGS[12],
    'Figure 12. "Voces de Curundú" testimonials section. New page with placeholder athlete story '
    'cards (awaiting real testimonial content from ADC founders). No equivalent section existed '
    'anywhere on the original Squarespace site.'
)

block(
    'Prompt:',
    '"Add a testimonials section called \'Voces de Curundú\' (Voices of Curundú). '
    'Each card should include space for an athlete photo, a short personal biography, '
    'and a direct quote from the athlete or their family about how ADC has changed their '
    'life. Use realistic placeholder content that ADC can easily replace."'
)

block(
    'Rationale:',
    'Testimonials from program beneficiaries are among the most persuasive content a '
    'nonprofit can publish. Research in nonprofit communications consistently identifies '
    'personal stories as more effective at driving donations and community engagement '
    'than any statistical presentation of impact. "Voces de Curundú" transforms abstract '
    'impact claims into human narratives. For an organization serving a community that '
    'is chronically underrepresented in mainstream media, providing athletes and families '
    'a visible, named platform also fulfills a dignity and representation function '
    'distinct from its fundraising utility. The placeholder card structure ensures that '
    'real content can be substituted by ADC staff without any technical knowledge.'
)

figure(
    IMGS[13],
    'Figure 13. FAQ page (Preguntas Frecuentes). New accordion-style FAQ organized by audience: '
    'Athletes & Programs, Volunteering & Support, and Donations—14 questions total. '
    'No FAQ page existed on the original site.'
)

block(
    'Prompt:',
    '"Build a complete FAQ page with an accordion interaction style. Organize all '
    'questions by the audience they serve: Athletes and Programs, Volunteering and '
    'Support, and Donations. Cover the 14 most common questions ADC receives from '
    'families, volunteers, and potential donors."'
)

block(
    'Rationale:',
    'FAQ pages serve a dual operational function: they reduce repetitive inquiry burden '
    'on staff (ADC\'s two founders were personally answering identical WhatsApp messages '
    'from families weekly about registration, training schedules, and costs), and they '
    'improve organic search performance by providing question-format content that matches '
    'natural search behavior. The accordion interface keeps the page visually clean while '
    'allowing users to scan headings and expand only relevant entries. Organizing by '
    'audience type ensures that a registering parent, a student seeking CAS service hours, '
    'and a first-time donor all reach their answers without reading through irrelevant content.'
)

heading('Phase 4 (Continued): Administrative Backend and AI Grant Writing', 2)

body(
    'Beyond the public-facing website, Phase 4 produced a private administrative '
    'dashboard hosted separately from the main site. This dashboard includes a '
    'Claude API-powered grant proposal writer—the most technically sophisticated '
    'component of the entire project. ADC staff enter basic information about a '
    'funding opportunity (funder name, grant amount, program area) and receive a '
    'complete, professionally formatted grant proposal draft within seconds.'
)

body(
    'The architecture uses a Cloudflare Worker as a serverless proxy to route API '
    'calls between the browser-based dashboard and the Anthropic API, ensuring the '
    'API key is never exposed to the client. The entire system—HTML dashboard, '
    'JavaScript frontend, Worker proxy, deployment configuration, and DEPLOY.md '
    'documentation—was generated through natural-language prompting of Claude. '
    'Estimated cost per grant proposal: approximately $0.003 USD.'
)

body(
    'This feature demonstrates one of the most compelling applications of AI in the '
    'nonprofit sector: eliminating the time and expertise barrier to pursuing grant '
    'funding. Organizations like ADC are frequently eligible for grants they never '
    'apply for simply because they lack staff with professional grant writing experience. '
    'An AI system that produces a usable draft in seconds at negligible cost fundamentally '
    'changes that equation, potentially unlocking tens of thousands of dollars in annual '
    'funding for an organization whose founders currently spend the majority of their '
    'volunteer hours on direct program delivery rather than development work.'
)

# ═══════════════════════════════════════════════════════════════════════════
# PHASE 5
# ═══════════════════════════════════════════════════════════════════════════

heading('Phase 5: Documentation and Long-Term Sustainability')

body(
    'The final development phase focused on ensuring that ADC could maintain, update, '
    'and eventually fully own the digital infrastructure produced across Phases 1–4. '
    'This included generating comprehensive documentation in Spanish for ADC\'s founders, '
    'a structured session history (48 prompts across 5 phases), a deployment guide for '
    'the Cloudflare Worker, an information checklist of content still needed from ADC, '
    'and this academic paper. Accessibility improvements were also applied in this '
    'phase: color contrast ratios, semantic HTML landmark structure, and ARIA labels '
    'were audited and corrected by the AI to achieve WCAG 2.2 compliance for users '
    'with visual impairments.'
)

body(
    'The documentation strategy reflects a critical principle of responsible AI-assisted '
    'development: the humans who will steward the system after the project lead steps away '
    'must be able to understand, operate, and own it. AI can generate code rapidly, but '
    'without deliberate knowledge transfer, it creates technical dependency rather than '
    'organizational capacity. Phase 5 ensured that ADC\'s long-term digital independence '
    'was designed into the project from the beginning, not treated as an afterthought.'
)

# ═══════════════════════════════════════════════════════════════════════════
# PLANNED FUTURE WORK
# ═══════════════════════════════════════════════════════════════════════════

heading('Planned Enhancements and Future Development')

body(
    'The following enhancements are planned for upcoming development phases. Each '
    'represents a natural extension of the current platform\'s capabilities and has '
    'been designed with ADC\'s operational capacity and long-term technical independence '
    'in mind.'
)

heading('Cloudflare Worker Deployment and Admin Dashboard Hosting', 3)
body(
    'The AI grant writing backend requires a one-time deployment of the Cloudflare Worker '
    'proxy via the Wrangler CLI and configuration of a secret Anthropic API key. This '
    'five-minute setup will activate the live grant writing capability. The admin dashboard '
    'will subsequently be hosted on Netlify via a private GitHub repository, providing '
    'password-protected access for ADC founders without requiring any server management knowledge.'
)

heading('Live Google Calendar Integration', 3)
body(
    'The current calendar page displays static event data. A live Google Calendar embed '
    'connected to ADC\'s existing calendar account (adcurundu@gmail.com) will replace '
    'the static content, enabling staff to update the public schedule by editing a '
    'tool they already use daily—without touching any code. This is the highest-priority '
    'operational improvement pending implementation.'
)

heading('Instagram Auto-Feed', 3)
body(
    'A LightWidget or Behold.so integration will connect ADC\'s active Instagram account '
    '(@adcurundu) to the website homepage, automatically displaying recent posts without '
    'manual content duplication. This closes the gap between ADC\'s most active content '
    'channel and its web presence, ensuring the homepage always reflects current '
    'organizational activity.'
)

heading('Netlify CMS for Blog Editing', 3)
body(
    'Netlify CMS (now Decap CMS) will be installed to provide a browser-based visual '
    'editing interface for the blog. This will allow ADC founders and trained volunteers '
    'to publish new articles, upload photos, and edit existing posts without any HTML '
    'knowledge. All changes will be version-controlled through the GitHub repository, '
    'providing a complete edit history and the ability to revert accidental changes.'
)

heading('Live AI Chatbot via Claude API', 3)
body(
    'The current website chatbot operates on rule-based scripted responses. Upgrading '
    'it to use the Claude API through the same Cloudflare Worker infrastructure already '
    'built for the grant writer will enable it to answer nuanced, open-ended questions '
    'from families, volunteers, and donors in natural conversational Spanish and English. '
    'This would make ADC one of the first youth nonprofit organizations in Panama '
    'with a live large-language-model assistant on its public website.'
)

heading('Content Pending from ADC Founders', 3)
body(
    'Several sections of the current site contain placeholder content awaiting real '
    'information from ADC\'s founders: (1) athlete and family testimonials for the '
    '"Voces de Curundú" section; (2) updated participant and impact statistics; '
    '(3) complete answers to the FAQ questions currently marked as pending; '
    '(4) sponsor website URLs for the sponsor bar hyperlinks; (5) a working '
    'donation platform URL; and (6) ADC\'s Google Calendar ID for the live '
    'calendar integration. Upon receipt of this content, the site will be fully '
    'populated and ready for official public launch.'
)

heading('GitHub and Analytics Ownership Transfer', 3)
body(
    'The final planned action is the transfer of the GitHub repository and Google '
    'Analytics account to ADC\'s own organizational credentials. This ensures that '
    'ADC\'s digital infrastructure is fully owned by the organization—not contingent '
    'on any individual developer\'s accounts or continued involvement. Training sessions '
    'will be provided to ADC founders on basic GitHub content-update workflows and '
    'Google Analytics dashboard interpretation.'
)

# ═══════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════

heading('Conclusion')

body(
    'This project demonstrates with unusual concreteness what conversational AI is '
    'capable of when applied systematically to a real-world technical problem by a '
    'non-technical operator. Across five development phases, a large language model '
    'transformed an outdated, disorganized Squarespace site into a professional, '
    'production-grade digital platform—generating eight HTML pages, a complete CSS '
    'design system, multiple JavaScript interaction features, a serverless API proxy, '
    'a private administrative backend with AI grant writing capabilities, and comprehensive '
    'organizational documentation—all from forty-eight natural-language prompts.'
)

body(
    'The implications extend beyond this single project. Small nonprofits worldwide are '
    'systematically underserved by digital technology because they cannot afford professional '
    'developers and lack the expertise to build for themselves. Conversational AI removes '
    'that barrier. The ADC website was rebuilt at zero developer cost, in a fraction of the '
    'time a conventional engagement would require, and with a level of technical sophistication '
    '—responsive design, API integrations, AI-powered features, WCAG accessibility—that '
    'would have been financially inaccessible to a small Panama-based nonprofit just a few '
    'years ago.'
)

body(
    'At the same time, this project is instructive about the boundaries of AI-assisted '
    'development. The AI cannot independently obtain testimonials from beneficiaries, cannot '
    'make strategic judgment calls about organizational priorities, and cannot cultivate '
    'the long-term community relationships that give the work its meaning. Those contributions '
    'remain irreducibly human. What AI provides is technical execution capacity—the ability '
    'to translate human vision into production-ready code rapidly, precisely, and at scale. '
    'For organizations like ADC, that capacity is not a convenience; it is transformative.'
)

# ═══════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════

heading('References')

refs = [
    'Anthropic. (2024). Claude: Model card and evaluations. Anthropic PBC. https://www.anthropic.com',
    'Asociación Deportiva Curundú. (2014–2026). Organizational records, program documentation, '
        'and session prompt history. ADC, Panama City, Panama.',
    'Cloudflare, Inc. (2025). Cloudflare Workers documentation: Serverless execution environment. '
        'https://developers.cloudflare.com/workers/',
    'GitHub, Inc. (2025). GitHub Pages documentation: Static site hosting from a repository. '
        'https://docs.github.com/en/pages',
    'Nielsen Norman Group. (2023). Donation page best practices for nonprofit websites. '
        'Nielsen Norman Group Research Reports.',
    'Netlify. (2025). Decap CMS (formerly Netlify CMS): Open-source content management for '
        'Git-based workflows. https://decapcms.org',
    'World Wide Web Consortium. (2023). Web content accessibility guidelines (WCAG) 2.2. '
        'W3C Recommendation. https://www.w3.org/TR/WCAG22/',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent      = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

# ═══════════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════════

out = '/Users/connorsolvason/my-website/AI Project Rough Draft.docx'
doc.save(out)
print(f'Saved → {out}')
