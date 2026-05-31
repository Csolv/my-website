"""
Converts AI Project Rough Draft.docx → ADC AI Project Final.docx
Changes:
  - Date updated to April 8, 2026
  - Prompt count updated: 48 → 56 throughout
  - Phase 5 extended with Prompts 52–56
  - New "Project Deliverables" section with showcase link + QR code
  - Conclusion updated
Run: python3 finalize_paper.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy, os

BASE  = "/Users/connorsolvason/my-website"
SRC   = f"{BASE}/AI Project Rough Draft.docx"
DST   = f"{BASE}/ADC AI Project Final.docx"
QR    = f"{BASE}/ADC_Code_Showcase_QR.png"
URL   = "https://csolv.github.io/my-website/ADC_Code_Showcase.html"

doc = Document(SRC)

# ── Helpers ───────────────────────────────────────────────────────────────────

def replace_text_in_run(run, old, new):
    if old in run.text:
        run.text = run.text.replace(old, new)

def replace_text_in_para(para, old, new):
    for run in para.runs:
        replace_text_in_run(run, old, new)
    # Also fix cross-run splits
    if old in para.text and old not in "".join(r.text for r in para.runs):
        # rebuild all text into first run
        full = para.text.replace(old, new)
        for i, r in enumerate(para.runs):
            r.text = full if i == 0 else ""


def _new_p_elem(bold_label=None, normal_text=None, bold=False, italic=False):
    """
    Returns a raw <w:p> element.
    If bold_label is provided, first run is bold; normal_text run follows.
    """
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Normal')
    pPr.append(pStyle)
    p.append(pPr)

    def make_run(text, b=False, i=False):
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if b:
            rPr.append(OxmlElement('w:b'))
        if i:
            rPr.append(OxmlElement('w:i'))
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        if text and (text[0] == ' ' or text[-1] == ' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p.append(r)
        return r

    if bold_label:
        make_run(bold_label, b=True)
    if normal_text:
        make_run(normal_text, b=False, i=italic)
    elif bold_label is None:
        make_run("", b=bold, i=italic)

    return p


def insert_after(anchor_p_elem, new_p_elem):
    """Insert new_p_elem immediately after anchor_p_elem; returns new_p_elem."""
    anchor_p_elem.addnext(new_p_elem)
    return new_p_elem


def insert_image_after(anchor_p_elem, img_path, width_inches=2.5):
    """Insert a paragraph containing an image after anchor_p_elem."""
    # Create a temporary paragraph at end, add image, then move
    temp_para = doc.add_paragraph()
    run = temp_para.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    # Move to right position
    p_elem = temp_para._element
    p_elem.getparent().remove(p_elem)
    anchor_p_elem.addnext(p_elem)
    return p_elem


def add_hyperlink(para_elem, text, url):
    """Add a hyperlink run to an existing paragraph element."""
    part = doc.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    hl.append(r)
    para_elem.append(hl)
    return hl


# ── 1. Update date ────────────────────────────────────────────────────────────
replace_text_in_para(doc.paragraphs[2], 'March 25, 2026', 'April 8, 2026')

# ── 2. Update abstract ────────────────────────────────────────────────────────
abs_para = doc.paragraphs[7]
replace_text_in_para(abs_para,
    'five development phases spanning approximately four months',
    'five development phases spanning approximately six months')
replace_text_in_para(abs_para,
    'forty-eight natural-language prompts',
    'fifty-six natural-language prompts')
replace_text_in_para(abs_para,
    'The paper concludes with a review of planned enhancements that illustrate the expanding and ongoing role of conversational AI in organizational digital strategy.',
    'The paper concludes with a review of planned enhancements and documents the final submission artifacts produced for this course: an interactive code exhibit published to GitHub Pages and a slide presentation. The code exhibit provides direct access to all 6,346 lines of AI-generated source code, organized by file with the prompts that produced them.')

# ── 3. Update methodology prompt count ───────────────────────────────────────
replace_text_in_para(doc.paragraphs[18],
    'Forty-eight prompts were documented across the five phases.',
    'Fifty-six prompts were documented across the five phases.')

# ── 4. Update Phase 5 opening line ───────────────────────────────────────────
replace_text_in_para(doc.paragraphs[147],
    '(48 prompts across 5 phases)',
    '(56 prompts across 5 phases)')

# ── 5. Update conclusion prompt count ────────────────────────────────────────
replace_text_in_para(doc.paragraphs[168],
    'all from forty-eight natural-language prompts',
    'all from fifty-six natural-language prompts')

# ── 6. Insert new Phase 5 content after paragraph 148 ────────────────────────
#
# Insert order: we chain each paragraph off the previous one.
# Start anchor = doc.paragraphs[148]._element

anchor = doc.paragraphs[148]._element

# --- Phase 5 continued heading ---
h = _new_p_elem(bold_label='Phase 5 (Continued): Final Documentation, Presentation, and Submission')
anchor = insert_after(anchor, h)

# --- Intro paragraph ---
intro = (
    'The final stage of Phase 5 produced the submission deliverables for this course project: '
    'an interactive code exhibit HTML file, a slide presentation, a QR code, and this finalized paper. '
    'Five additional prompts were documented during this stage, bringing the total to 56 across '
    'the five phases.'
)
p = _new_p_elem(normal_text=intro)
anchor = insert_after(anchor, p)

# --- Prompt 52 ---
p = _new_p_elem(bold_label='Prompt 52 (April 7, 2026): Update prompt history and course paper')
anchor = insert_after(anchor, p)

p = _new_p_elem(bold_label='Request:  ',
                normal_text='Update the ADC prompt history Word document and the course paper rough '
                            'draft with all prompts documented since the last session, including '
                            'Phase 4 prompts 42\u201351 and the Cloudflare Worker deployment.')
anchor = insert_after(anchor, p)

p = _new_p_elem(bold_label='Result:  ',
                normal_text='ADC_Prompts_Historial.docx updated \u2014 Phase 4 header added, '
                            'Prompts 42\u201351 documented with Request and Result for each session. '
                            'AI Project Rough Draft.docx updated \u2014 \u201cCloudflare Worker Deployment\u201d '
                            'section marked [COMPLETE] and Worker live URL added.')
anchor = insert_after(anchor, p)

# --- Prompt 53 ---
p = _new_p_elem(bold_label='Prompt 53 (April 8, 2026): Create 3-minute FECON 390 presentation')
anchor = insert_after(anchor, p)

p = _new_p_elem(bold_label='Request:  ',
                normal_text='Using the course paper as a foundation, create a 3-minute slides '
                            'presentation for FECON 390. Include key prompts, side-by-side website '
                            'comparisons using the 13 screenshots, and key project details. Tailor '
                            'it to show the power of AI \u2014 emphasize change, not process.')
anchor = insert_after(anchor, p)

p = _new_p_elem(bold_label='Result:  ',
                normal_text='ADC_AI_Presentation.pptx created \u2014 12 slides in .pptx format '
                            '(Google Slides-importable). Includes title, problem setup, method overview, '
                            '3 side-by-side screenshot comparison slides, actual prompt quotes, '
                            'new-pages showcase, AI grant writer slide, project statistics panel, '
                            'and Q&A slide with live site links.')
anchor = insert_after(anchor, p)

# --- Prompt 54 ---
p = _new_p_elem(bold_label='Prompt 54 (April 8, 2026): Build HTML code showcase for project submission')
anchor = insert_after(anchor, p)

p = _new_p_elem(bold_label='Request:  ',
                normal_text='Create a browser-openable HTML code showcase for FECON 390 project '
                            'submission. Include a full file tree, line counts, and annotated code '
                            'snippets from each file \u2014 each labeled with the prompt that generated it. '
                            'Redact the password hash and live Worker URL before publishing.')
anchor = insert_after(anchor, p)

p = _new_p_elem(bold_label='Result:  ',
                normal_text='ADC_Code_Showcase.html created and deployed to GitHub Pages. '
                            'Documents all 13 source files (6,346 lines), 6 annotated snippets '
                            'with Prism.js syntax highlighting, and prompt-to-code mapping for '
                            'each file. SHA-256 hash and Worker URL redacted. Live at: '
                            + URL)
anchor = insert_after(anchor, p)

# --- Prompt 55 ---
p = _new_p_elem(bold_label='Prompt 55 (April 8, 2026): Generate QR code for code showcase')
anchor = insert_after(anchor, p)

p = _new_p_elem(bold_label='Request:  ',
                normal_text='Generate a QR code that links directly to the live code showcase on GitHub Pages.')
anchor = insert_after(anchor, p)

p = _new_p_elem(bold_label='Result:  ',
                normal_text='ADC_Code_Showcase_QR.png created \u2014 ADC green on white, '
                            'high error correction (H level), 564\u00d7564px. '
                            'Links to the live code exhibit URL.')
anchor = insert_after(anchor, p)

# --- Prompt 56 ---
p = _new_p_elem(bold_label='Prompt 56 (April 8, 2026): Finalize course paper')
anchor = insert_after(anchor, p)

p = _new_p_elem(bold_label='Request:  ',
                normal_text='Update the AI project rough draft to its final version. Update the date '
                            'and prompt counts throughout, add the five most recent prompts, add a '
                            'project deliverables section with the HTML showcase link and QR code, '
                            'and save as the final paper.')
anchor = insert_after(anchor, p)

p = _new_p_elem(bold_label='Result:  ',
                normal_text='AI Project Rough Draft.docx \u2192 ADC AI Project Final.docx. '
                            'Date updated to April 8, 2026. Prompt count updated to 56 throughout. '
                            'Phase 5 continued section added with Prompts 52\u201356. '
                            'Project deliverables section added with code showcase link and QR code. '
                            'Abstract and conclusion updated.')
anchor = insert_after(anchor, p)

# --- Project Deliverables section heading ---
empty = _new_p_elem(normal_text='')
anchor = insert_after(anchor, empty)

h2 = _new_p_elem(bold_label='Project Deliverables')
anchor = insert_after(anchor, h2)

# --- Deliverables intro paragraph ---
deliv_intro = (
    'Three artifacts were produced as final submission deliverables for FECON 390. '
    'Each is a direct output of AI-assisted development and is designed to be '
    'accessible to reviewers, classmates, and anyone interested in the project.'
)
p = _new_p_elem(normal_text=deliv_intro)
anchor = insert_after(anchor, p)

# --- Code Exhibit heading ---
p = _new_p_elem(bold_label='     Interactive Code Exhibit', italic=True)
anchor = insert_after(anchor, p)

# --- Code exhibit description paragraph with hyperlink ---
desc_p_elem = _new_p_elem(
    normal_text='The complete project source code is published as an interactive HTML exhibit at: '
)
# Add the hyperlink inline
add_hyperlink(desc_p_elem, URL, URL)
# Add closing text
r_close = OxmlElement('w:r')
t_close = OxmlElement('w:t')
t_close.text = ('. The exhibit documents all 13 source files (6,346 total lines of AI-generated code), '
                'organized by file with descriptions, the specific prompt that generated each file, '
                'and annotated code snippets with full syntax highlighting. It is accessible via '
                'the QR code below or the URL above.')
t_close.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
r_close.append(t_close)
desc_p_elem.append(r_close)
anchor = insert_after(anchor, desc_p_elem)

# --- QR code image ---
if os.path.exists(QR):
    qr_anchor = insert_image_after(anchor, QR, width_inches=2.2)
    anchor = qr_anchor
    # Caption
    cap_p = _new_p_elem(normal_text='Figure 14. QR code linking to the live code exhibit at csolv.github.io/my-website/ADC_Code_Showcase.html')
    # Make caption italic small
    for r in cap_p.iter(qn('w:r')):
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        i_elem = OxmlElement('w:i')
        rPr.append(i_elem)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '18')  # 9pt
        rPr.append(sz)
    anchor = insert_after(anchor, cap_p)

# --- Slide Presentation entry ---
p = _new_p_elem(bold_label='     Slide Presentation', italic=True)
anchor = insert_after(anchor, p)

p = _new_p_elem(normal_text=(
    'A 12-slide PowerPoint presentation (.pptx) summarizes the project for FECON 390. '
    'It includes three side-by-side before-and-after website comparisons using the '
    '13 screenshots captured at project completion, actual prompt quotes, a statistical '
    'overview of the project scope, and a summary of the AI grant writing backend. '
    'The file (ADC_AI_Presentation.pptx) is included in the project submission folder '
    'and is importable to Google Slides.'
))
anchor = insert_after(anchor, p)

# --- blank separator ---
p = _new_p_elem(normal_text='')
anchor = insert_after(anchor, p)

# ── 7. Save ───────────────────────────────────────────────────────────────────
doc.save(DST)
print(f'Saved: {DST}')
print(f'Total paragraphs: {len(doc.paragraphs)}')
