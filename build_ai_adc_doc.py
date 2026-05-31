"""
Creates ADC_AI_Integrations.docx in /my-website
Run: python3 build_ai_adc_doc.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

GREEN      = RGBColor(0x1B, 0x6B, 0x3A)
GREEN_LIGHT= RGBColor(0x2E, 0xAD, 0x6B)
DARK       = RGBColor(0x0D, 0x0D, 0x0D)
GRAY       = RGBColor(0x55, 0x55, 0x55)

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=GREEN)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_font(run, size=12, bold=True, color=DARK)
    return p

def subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, italic=True, color=GREEN_LIGHT)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, size=11, color=DARK)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11, color=DARK)
    return p

def label_line(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(label + '  ')
    set_font(r1, size=10, bold=True, color=GREEN)
    r2 = p.add_run(text)
    set_font(r2, size=10, color=GRAY)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2EAD6B')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run('AI with ADC')
set_font(r, size=26, bold=True, color=GREEN)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Future Technology Integrations for Asociación Deportiva Curundú')
set_font(r, size=13, italic=True, color=GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run('Connor Solvason  ·  April 2026')
set_font(r, size=10, color=GRAY)

divider()

# ── Design Principles ────────────────────────────────────────────────────────
heading1('Design Principles — Non-Negotiables for All Features')

principles = [
    ('$0 cost',             'Every tool must be free to use. No subscriptions, no API fees, no hidden costs.'),
    ('No coding required',  'ADC founders should interact with every feature the way they use Squarespace — drag, click, and fill in fields. No prompts, no terminal, no code.'),
    ('Intuitive backend',   'If a founder needs instructions to use a feature, it is not finished. The UI must be self-explanatory.'),
    ('Squarespace-compatible', 'All public-facing features must be embeddable in a Squarespace site via iframe or embed block so ADC can host everything from one platform.'),
    ('Founder-first',       'Designed for two busy nonprofit founders with no technical staff. Every feature should save time, not create new tasks.'),
]

for label, desc in principles:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run('▸  ' + label + ': ')
    set_font(r1, size=11, bold=True, color=GREEN)
    r2 = p.add_run(desc)
    set_font(r2, size=11, color=DARK)

divider()

# ════════════════════════════════════════════════════════════════════════════
# FEATURES
# ════════════════════════════════════════════════════════════════════════════

# ── 1. Grant Roadmap ─────────────────────────────────────────────────────────
heading1('1.  Interactive Grant Roadmap for Panama')
subheading('Admin page feature — internal use by founders')

body('A visual, road-style map that walks ADC founders through every viable grant source available to a Panama-based nonprofit. The roadmap is interactive: each grant is a stop on the road with an expandable checklist, friendly tip callout, and timeline estimate.')

heading2('What founders see')
bullet('A scrollable road graphic with grant "stops" (e.g., USAID, Inter-American Foundation, Duke grants, Panama government, private foundations)')
bullet('Each stop expands to show: eligibility checklist, required documents, typical award range, application deadline cycle, and a friendly tip from past applicants')
bullet('Green checkmark when a step is complete — progress saves automatically to the browser')
bullet('Color-coded timeline bar showing which grants are open now vs. in 3/6/12 months')
bullet('Filter by: grant size, application complexity, likelihood for first-time applicants')

heading2('Implementation')
label_line('Frontend:', 'Vanilla HTML/CSS/JS — scrollable SVG road graphic, accordion checklist per stop')
label_line('Storage:', 'localStorage saves checklist state per grant — no database, no login required beyond existing admin panel')
label_line('Data:', 'Grant data maintained as a JSON object founders can edit via the admin panel (add/remove grants, update deadlines)')
label_line('Cost:', '$0 — fully static, no external APIs')
label_line('Squarespace:', 'Embed via iframe on a password-protected Squarespace page for internal use')

divider()

# ── 2. Duke Alum Map ─────────────────────────────────────────────────────────
heading1('2.  Live Duke Alumni Source Map — Panama')
subheading('Admin page feature — donor development tool')

body('An interactive map showing Duke University alumni located in Panama. Each pin on the map opens a profile card with a photo, brief bio, their organization, and any known history of philanthropy or connection to ADC. Designed to help founders identify and warm-contact potential major donors and partners.')

heading2('What founders see')
bullet('A map of Panama with colored pins for each Duke alum')
bullet('Click a pin → profile card slides open with: name, photo, title/organization, Duke graduation year, philanthropic history, ADC connection notes')
bullet('Filter panel: by industry, by location, by giving history, by connection strength')
bullet('Add new alum via a simple form — fill in fields, upload photo, hit Save')
bullet('Export to CSV for outreach planning')

heading2('Implementation')
label_line('Map:', 'Leaflet.js (free, open-source) with OpenStreetMap tiles — no Google Maps API cost')
label_line('Data:', 'JSON file edited via admin panel form — founders add/edit alumni without touching code')
label_line('Photos:', 'Stored as file uploads or as LinkedIn photo URLs')
label_line('Source data:', 'Pre-populated from Duke Alumni Association (public profiles), LinkedIn, DukeEngage contacts, and existing ADC network')
label_line('Cost:', '$0 — Leaflet + OpenStreetMap are fully free')
label_line('Squarespace:', 'Embed Leaflet map as iframe on admin or partner page')

divider()

# ── 3. Panama Legal Guide ─────────────────────────────────────────────────────
heading1('3.  Panama 501(C) Legal & Admin Guide')
subheading('Admin page feature — dynamic reference tool')

body('A structured, plain-language guide to the legal and administrative requirements for operating a nonprofit in Panama. Organized by topic with money-saving tips, deadline reminders, and links to official sources. Founders can mark items as done and add their own notes.')

heading2('What founders see')
bullet('Accordion-style guide organized by section: Registration & Legal Status, Tax Exemptions, Banking Requirements, Annual Reporting, Fundraising Rules, International Funding Compliance')
bullet('Each section has: plain-language summary, official requirement, money-saving tip, common mistake warning, and link to official document')
bullet('Editable notes field per section so founders can record their specific situation')
bullet('Checklist mode: mark requirements as completed, pending, or not applicable')
bullet('Deadline calendar: annual filings and renewals shown on a visual timeline')

heading2('Key content areas')
bullet('Fundación vs. Asociación structure — which is better for ADC and why')
bullet('Ministerio de Gobierno registration process and renewal requirements')
bullet('ITBMS (Panama VAT) exemption status for nonprofits')
bullet('Foreign donation reporting under Panama AML/CFT regulations')
bullet('Panama Foundation Law (Law 25 of 1995) — key provisions for ADC')
bullet('DGI (tax authority) filing requirements and penalties')
bullet('Money-saving tip: which legal steps can be done without an attorney')

heading2('Implementation')
label_line('Frontend:', 'Static HTML/JS accordion — all content stored in a JSON file founders can update')
label_line('Updates:', 'Admin form lets founders add new items, edit descriptions, or mark laws as changed')
label_line('Cost:', '$0 — fully static, no external dependencies')
label_line('Squarespace:', 'Embed on password-protected admin page via iframe')

divider()

# ── 4. Social Media Campaign Generator ───────────────────────────────────────
heading1('4.  Social Media Campaign Generator')
subheading('Admin page feature — content creation tool')

body('A drag-and-fill template generator for ADC social media posts. Founders pick a template, drop in photos, and download a ready-to-post graphic. Templates are based on past ADC posts. Optional: type a creative prompt (e.g., "kids as superheroes") to generate a themed template suggestion.')

heading2('What founders see')
bullet('Gallery of pre-built templates: match-day post, fundraising ask, milestone celebration, player spotlight, event announcement, sponsor thank-you')
bullet('Click a template → drag to reorder elements, click text to edit, click image placeholder to upload photo')
bullet('History panel: past campaigns saved automatically so founders can re-use or update previous templates')
bullet('Creative prompt field: type "kids as pro players" → AI suggests layout variation or overlay style matching the theme')
bullet('Export as PNG/JPG at Instagram and Facebook resolution in one click')

heading2('Implementation')
label_line('Canvas editor:', 'Fabric.js (free, open-source) — browser-based drag-and-drop canvas with text, image, and overlay layers')
label_line('Templates:', 'Pre-built as Fabric.js JSON configs; founders cannot break them — only fill in content')
label_line('History:', 'Saved to localStorage as canvas JSON snapshots')
label_line('AI prompts:', 'Optional: calls existing Cloudflare Worker (already deployed) to generate layout suggestions — no new API cost')
label_line('Cost:', '$0 — Fabric.js is free; reuses existing Worker')
label_line('Squarespace:', 'Embed on admin page iframe; output images downloaded to device for posting directly to Instagram/Facebook')

divider()

# ── 5. Documentary ────────────────────────────────────────────────────────────
heading1('5.  ADC Full Documentary')
subheading('Public-facing — YouTube / website embed')

body('A short (8–15 min) documentary film about ADC\'s founding, community, athletes, and impact. Designed to serve as the single most powerful fundraising and grant-application tool — a video any donor or grant reviewer can watch to immediately understand who ADC is and why it matters.')

heading2('Content outline')
bullet('Opening: aerial shot of Curundú neighborhood, narrated context about the community')
bullet('Chapter 1 — The Founders: interviews with Andrés Madrid, César Santos, and Javier Wallace; their story, why they started ADC in 2014')
bullet('Chapter 2 — The Athletes: follow 2–3 players through a week of training, school, and life in Curundú; let them speak')
bullet('Chapter 3 — The Program: training sessions, nutrition meals, tutoring — show, don\'t tell')
bullet('Chapter 4 — The Impact: players who reached national teams; families who stayed; the alternative without ADC')
bullet('Closing: call to action — donate, volunteer, partner')

heading2('Production approach')
bullet('Shot on iPhone with a DJI Osmo Mobile gimbal ($0 if borrowed) — modern smartphone footage is documentary-quality')
bullet('Edit in CapCut or DaVinci Resolve (both free)')
bullet('Score with royalty-free music from Pixabay or YouTube Audio Library ($0)')
bullet('Subtitles in both Spanish and English for grant reviewer audiences')

label_line('Hosting:', 'YouTube (free, unlimited) — embed on ADC website and include link in all grant applications')
label_line('Cost:', '$0 if shot on existing devices; minimal if renting gimbal')
label_line('Squarespace:', 'YouTube embed block — native Squarespace feature, no code needed')

divider()

# ── 6. 2D Interactive Campus View ────────────────────────────────────────────
heading1('6.  2D Interactive ADC Campus — Google Earth Zoom View')
subheading('Public-facing website feature')

body('A top-down 2D illustrated view of a typical ADC day — styled like a Google Earth zoom from neighborhood level down to street level and then into each activity zone. Visitors click on zones to zoom in and see what happens there, with photos, video clips, and donation links embedded in each zone.')

heading2('Zone breakdown')
bullet('Soccer Field — the field: shows training drills in action, trophy showcase, equipment used; donate-a-ball / donate-a-uniform button embedded')
bullet('Sidelines / Nutrition Station: shows meal preparation, what the nutrition program provides, how families can contribute food donations')
bullet('Education Corner: students doing homework, tutors helping; shows textbooks and supplies; donate-a-backpack link')
bullet('Community Entrance: map of Curundú neighborhood with context about who lives there; connection to the "why"')
bullet('Trophy Room (overlay): click to see all tournament wins and national team players')

heading2('Implementation')
label_line('Illustration:', 'Custom SVG map with clickable zones — built once, reusable forever')
label_line('Zoom animation:', 'CSS transform + JS — smooth zoom-in effect on click, no video required')
label_line('Media:', 'Each zone panel pulls from existing ADC photos and videos already on the website')
label_line('Cost:', '$0 — pure HTML/CSS/JS + existing media assets')
label_line('Squarespace:', 'Embed as iframe on homepage or dedicated "Our World" page')

divider()

# ── 7. Offline Mini Game ─────────────────────────────────────────────────────
heading1('7.  Offline Soccer Mini Game — "ADC Run"')
subheading('Public-facing — playable by kids on the website')

body('A browser-based endless runner game — like the Chrome Dino game — but starring an ADC soccer player. The player dribbles a ball and must jump over sliding tackles or duck under headers while keeping possession. Works offline using a Service Worker so kids can play even without internet access.')

heading2('Gameplay')
bullet('ADC-branded player sprite dribbling a soccer ball — can use real player silhouettes or cartoon versions')
bullet('Obstacles: sliding tackles (jump over), header balls (duck under), referee with red card (dodge)')
bullet('Power-ups: ADC trophy (score multiplier), nutrition meal (speed boost), teammate (shield from one tackle)')
bullet('High score saved to localStorage — kids compete for best score')
bullet('ADC green color scheme, simple chiptune sound effects')

heading2('Implementation')
label_line('Engine:', 'Vanilla JS canvas — no library needed, ~300 lines of code')
label_line('Offline:', 'Service Worker caches the game so it runs without internet once loaded once')
label_line('Assets:', 'Simple pixel sprites — can be designed free in Piskel (browser-based pixel editor)')
label_line('Cost:', '$0 — entirely self-contained')
label_line('Squarespace:', 'Embed as iframe; the game self-caches on first load so kids can play offline after that')

divider()

# ── 8. Fundraiser Photo Generator ────────────────────────────────────────────
heading1('8.  Fundraiser Concept Photo Generator')
subheading('Admin page feature — campaign visualization tool')

body('A tool that lets founders visualize what a fundraising project would look like before it exists — generating concept images for grant proposals, donor pitches, and social posts. Type a description, click Generate, get a realistic mockup image.')

heading2('What founders see')
bullet('Text field: "New practice jerseys with ADC logo for U14 team" → generates a realistic image of kids in the described jerseys')
bullet('Text field: "Covered outdoor nutrition station with tables for 30" → generates architectural-style concept image')
bullet('Gallery of past generated images saved by project name for reuse in proposals')
bullet('One-click export at print resolution for grant application attachments')

heading2('Implementation')
label_line('AI image gen:', 'Stable Diffusion via Hugging Face free inference API — $0 cost for reasonable usage')
label_line('Backend:', 'Extend existing Cloudflare Worker to route image generation requests — no new infrastructure')
label_line('Gallery:', 'Images stored as base64 in localStorage or as blob URLs')
label_line('Cost:', '$0 — Hugging Face free tier covers nonprofit-scale usage')
label_line('Squarespace:', 'Admin page only (iframe) — output images downloaded and inserted into grant docs manually')

divider()

# ── 9. ADC Theme Song ────────────────────────────────────────────────────────
heading1('9.  ADC Anthem / Theme Song')
subheading('Brand asset — website, video, social media')

body('A short (30–60 sec), catchy instrumental or vocal theme for ADC — used as the intro to the documentary, background for social media videos, website audio, and any future broadcast content. Gives ADC a recognizable sonic identity.')

heading2('Style direction')
bullet('Upbeat, energetic, Latin-influenced — reflects Panama and soccer culture')
bullet('Should feel like a youth sports anthem: inspiring, community-focused, not corporate')
bullet('Lyrics (if vocal): simple Spanish hook that kids would chant, centered on "Curundú," "nuestro equipo," or "ADC"')
bullet('Instrumentation: percussion-forward with guitar and brass; builds in energy')

heading2('Production approach')
bullet('Generate a reference track using Suno AI (free tier — generates full songs from text description in seconds)')
bullet('Prompt: "Upbeat Latin youth soccer anthem, Panama, inspirational, percussion and brass, builds to climax, 45 seconds, in Spanish"')
bullet('If full production desired: record with local musicians in Panama using the AI track as a reference — cost varies')
bullet('License: Suno free tier output is royalty-free for nonprofit use')

label_line('Cost:', '$0 for AI-generated version; minimal for live recording if desired')
label_line('Delivery:', 'MP3 file embedded on website, used in documentary, shared as Instagram Reel audio')
label_line('Tool:', 'suno.com — free, browser-based, no account required for basic use')

divider()

# ── Summary Table ─────────────────────────────────────────────────────────────
heading1('Summary — All Nine Features')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Feature overview at a glance:')
set_font(r, size=11, bold=True, color=DARK)

table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

hdr = table.rows[0].cells
headers = ['Feature', 'Location', 'Who uses it', 'Cost', 'Squarespace-ready']
for i, h in enumerate(headers):
    hdr[i].text = h
    for run in hdr[i].paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _el
    tc = hdr[i]._element
    tcPr = tc.get_or_add_tcPr()
    shd = _el('w:shd')
    shd.set(_qn('w:fill'), '1B6B3A')
    shd.set(_qn('w:color'), 'auto')
    shd.set(_qn('w:val'), 'clear')
    tcPr.append(shd)

rows_data = [
    ('Grant Roadmap',               'Admin page',    'Founders',          '$0', 'iframe embed'),
    ('Duke Alum Map',               'Admin page',    'Founders',          '$0', 'iframe embed'),
    ('Panama Legal Guide',          'Admin page',    'Founders',          '$0', 'iframe embed'),
    ('Social Media Generator',      'Admin page',    'Founders',          '$0', 'iframe embed'),
    ('Documentary',                 'YouTube/site',  'Public + donors',   '$0', 'YouTube block'),
    ('2D Interactive Campus',       'Public website','All visitors',       '$0', 'iframe embed'),
    ('Offline Soccer Game',         'Public website','Kids',               '$0', 'iframe embed'),
    ('Fundraiser Photo Generator',  'Admin page',    'Founders',          '$0', 'iframe embed'),
    ('ADC Theme Song',              'All media',     'All audiences',      '$0', 'audio embed'),
]

for row_data in rows_data:
    row = table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        for run in row[i].paragraphs[0].runs:
            run.font.size = Pt(9)

divider()

# ── Footer note ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r = p.add_run('All features are designed to be implemented incrementally — each is independent and can be built without the others being in place first. Priority order should be determined by ADC founders based on what saves the most time and creates the most immediate fundraising impact.')
set_font(r, size=10, italic=True, color=GRAY)

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(4)
r2 = p2.add_run('Connor Solvason  ·  FECON 390  ·  April 2026  ·  adcurundu.org')
set_font(r2, size=9, color=GRAY)

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/connorsolvason/my-website/ADC_AI_Integrations.docx'
doc.save(out)
print(f'Saved: {out}')
print(f'Paragraphs: {len(doc.paragraphs)}')
